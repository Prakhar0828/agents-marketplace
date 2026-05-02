"""AI Resume Optimizer.

Flow:
  1. User uploads a resume PDF (via the HTTP upload route) and types a prompt
     like "Senior Backend Engineer at Stripe".
  2. We pull job title + company out of the message with GPT. If the message
     or file is missing pieces, we ask a clarifying question in chat.
  3. Scrape LinkedIn jobs for that title + company via the Apify actor
     `bebity/linkedin-jobs-scraper` and pick the best-matching description.
  4. GPT rewrites the resume in markdown, tailored to the job description.
  5. We save the output as `.md` and `.docx` in DOWNLOADS_DIR and stream a
     `resume_result` event with both download URLs.
"""

from __future__ import annotations

import json
import re
import textwrap
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Awaitable, Callable

from docx import Document
from docx.shared import Pt
from mcp.shared.exceptions import McpError
from openai import OpenAI
from pypdf import PdfReader

from ..config import DOWNLOADS_DIR, MODEL
from ..mcp_client import LenientClientSession, call_tool
from ..routes.uploads import load_upload

Emit = Callable[[dict], Awaitable[None]]


# ── Data ──────────────────────────────────────────────────────────────────────


@dataclass
class Intent:
    job_title: str | None
    company: str | None
    location: str | None  # optional, improves LinkedIn match accuracy


# ── Intent extraction ─────────────────────────────────────────────────────────


def _extract_intent(
    client: OpenAI, user_message: str, prior: dict | None
) -> Intent:
    system = (
        "Extract resume-targeting intent from the user's message. "
        "Return JSON with keys:\n"
        "  job_title (str | null)  — the role they want to apply to\n"
        "  company   (str | null)  — the target company name\n"
        "  location  (str | null)  — optional city/country hint"
    )
    if prior:
        system += (
            "\n\nThe user was previously asked for missing info. "
            f"Merge their response into this prior intent: {json.dumps(prior)}. "
            "Keep existing non-null fields unless explicitly overridden."
        )

    resp = client.chat.completions.create(
        model=MODEL,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user_message},
        ],
    )
    data = json.loads(resp.choices[0].message.content)
    return Intent(
        job_title=data.get("job_title"),
        company=data.get("company"),
        location=data.get("location"),
    )


# ── PDF parsing ──────────────────────────────────────────────────────────────


def _parse_resume_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — malformed page shouldn't kill the flow
            continue
    text = "\n\n".join(p.strip() for p in pages if p.strip())
    # Collapse runs of whitespace that pypdf tends to leave behind.
    return re.sub(r"[ \t]+\n", "\n", text).strip()


# ── LinkedIn scrape ──────────────────────────────────────────────────────────


def _extract_json(text: str) -> list | dict | None:
    """Pull JSON out of a possibly-noisy MCP text blob.

    Apify's MCP shim typically streams two TextContent parts per call: the
    actual JSON result followed by a human-readable summary like
    `Actor "x" completed successfully! Run ID: …`. `call_tool` joins those
    with `\\n`, which is why a naive `json.loads` fails on `Extra data`.

    We use `raw_decode` to peel off every JSON value in order and flatten any
    top-level lists — so multiple streamed chunks still collapse into a single
    list of items.
    """
    if not text:
        return None

    decoder = json.JSONDecoder()
    values: list[Any] = []
    idx = 0
    n = len(text)
    while idx < n:
        # Skip whitespace and any leading non-JSON prose up to the next
        # structural character. raw_decode is strict about position, so we
        # scan forward until we hit a `{` or `[`.
        while idx < n and text[idx] not in "{[":
            idx += 1
        if idx >= n:
            break
        try:
            value, end = decoder.raw_decode(text, idx)
        except json.JSONDecodeError:
            # Not a real JSON value at this bracket — step past it and keep
            # looking. This handles the `[bracketed phrase]` inside a trailing
            # status line gracefully.
            idx += 1
            continue
        values.append(value)
        idx = end

    if not values:
        return None
    if len(values) == 1:
        return values[0]
    # Multiple values: flatten lists, otherwise return them as a list so the
    # caller decides.
    flat: list[Any] = []
    for v in values:
        if isinstance(v, list):
            flat.extend(v)
        else:
            flat.append(v)
    return flat


async def _scrape_linkedin_job(
    session: LenientClientSession,
    intent: Intent,
    emit: Emit,
) -> str | None:
    """Return the most relevant full job description, or None if nothing matched."""
    assert intent.job_title and intent.company  # narrowed before call

    # bebity/linkedin-jobs-scraper requires `location` — when the user didn't
    # give one we fall back to "Worldwide" which matches the actor's "any
    # region" sentinel and keeps the call valid.
    location = intent.location or "Worldwide"
    search_term = f"{intent.job_title} {intent.company}"
    # Note: deliberately omit `publishedAt` — the actor treats it as an enum
    # ("", "r86400", "r604800", "r2592000") and rejects anything else. Omitting
    # means "any time", which is what we want.
    inputs: dict[str, Any] = {
        "title": intent.job_title,
        "location": location,
        "companyName": [intent.company],
        "searchTerm": search_term,
        "rows": 10,
    }

    await emit(
        {
            "type": "tool_call",
            "name": "bebity--linkedin-jobs-scraper",
            "args": {
                "title": intent.job_title,
                "company": intent.company,
                "location": location,
            },
        }
    )

    try:
        raw = await call_tool(session, "bebity--linkedin-jobs-scraper", inputs)
    except McpError as exc:
        # Most common cause: the actor rejects our args, or the user's Apify
        # account hasn't enabled this actor yet. Surface a readable reason and
        # let the caller fall back to a generic JD.
        await emit(
            {
                "type": "tool_result",
                "name": "bebity--linkedin-jobs-scraper",
                "preview": f"Actor error: {exc}",
            }
        )
        return None

    data = _extract_json(raw)
    listings = data if isinstance(data, list) else []

    await emit(
        {
            "type": "tool_result",
            "name": "bebity--linkedin-jobs-scraper",
            "preview": f"{len(listings)} listing(s) returned",
        }
    )
    if not listings:
        return None

    # Prefer exact-company matches with a real description, newest first.
    company_lc = intent.company.lower()
    title_lc = intent.job_title.lower()

    def _rank(item: dict) -> tuple[int, int]:
        c = (item.get("companyName") or "").lower()
        t = (item.get("title") or "").lower()
        company_score = 2 if c == company_lc else (1 if company_lc in c else 0)
        title_score = 2 if title_lc in t else (1 if any(w in t for w in title_lc.split()) else 0)
        return (company_score + title_score, len(item.get("description") or ""))

    best = max(listings, key=_rank)
    description = best.get("description") or best.get("descriptionText") or ""
    return description.strip() or None


# ── Resume rewrite ───────────────────────────────────────────────────────────


RESUME_SYSTEM_PROMPT = textwrap.dedent(
    """\
    You are an expert resume writer and career coach.

    You will be given:
      1. The user's existing resume (raw text extracted from their PDF).
      2. A job description they're targeting.
      3. The target job title and company.

    Rewrite the resume to maximise relevance to that specific job. Rules:
      • Preserve all real, factual information (names, companies, dates, education).
        NEVER invent jobs, certifications, or skills the candidate doesn't have.
      • Mirror the language and keywords of the job description where truthful.
      • Reorder, regroup, and rephrase bullets to emphasise the most relevant
        experience first. Keep bullets action-verb-led and quantified when
        possible (reuse numbers/metrics already present in the original).
      • Add a concise Professional Summary (2-3 lines) tailored to this role.
      • Keep it to one page of content — prefer cutting less-relevant detail
        over cramming everything in.
      • Output valid GitHub-flavoured Markdown with this structure:

    # <Full Name>
    <contact line: email · phone · location · linkedin>

    ## Professional Summary
    ...

    ## Experience
    ### <Role> — <Company>
    *<Dates>*
    - bullet
    - bullet

    ## Skills
    - ...

    ## Education
    ### <Degree> — <Institution>
    *<Dates>*

    Also return a short (1-2 sentence) `summary` of the key changes you made.

    Respond with JSON of shape:
      {"markdown": "<full markdown resume>",
       "summary": "<what changed and why>"}
    """
)


def _rewrite_resume(
    client: OpenAI,
    resume_text: str,
    job_description: str,
    intent: Intent,
) -> tuple[str, str]:
    user_payload = {
        "target_job_title": intent.job_title,
        "target_company": intent.company,
        "job_description": job_description[:6000],  # cap context to protect budget
        "resume_text": resume_text[:12000],
    }
    resp = client.chat.completions.create(
        model=MODEL,
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": RESUME_SYSTEM_PROMPT},
            {"role": "user", "content": json.dumps(user_payload)},
        ],
    )
    data = json.loads(resp.choices[0].message.content)
    return data.get("markdown", "").strip(), data.get("summary", "").strip()


# ── Output generation ────────────────────────────────────────────────────────


def _slugify(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")
    return slug[:48] or "resume"


def _markdown_to_docx(markdown: str, path: Path) -> None:
    """Render the markdown resume into a tidy docx file.

    We don't try to be a full Markdown renderer — resumes have a small, known
    structure: `#` for name, `##` for section, `###` for role header, `*text*`
    for italic date line, and `- ` for bullets. Anything else we drop in as a
    plain paragraph.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("*") and line.endswith("*") and len(line) > 2:
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True
        else:
            doc.add_paragraph(line)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _save_outputs(markdown: str, intent: Intent) -> tuple[str, str]:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    slug = f"{_slugify(intent.company or 'target')}_{_slugify(intent.job_title or 'role')}"
    stem = f"resume_{slug}_{timestamp}"

    md_path = DOWNLOADS_DIR / f"{stem}.md"
    md_path.write_text(markdown, encoding="utf-8")

    docx_path = DOWNLOADS_DIR / f"{stem}.docx"
    _markdown_to_docx(markdown, docx_path)

    return md_path.name, docx_path.name


# ── Public entrypoint ─────────────────────────────────────────────────────────


ResumeState = dict[str, Any]


def new_state() -> ResumeState:
    return {"intent": None, "awaiting": None, "resume_text": None, "resume_name": None}


async def run_resume_agent(
    user_message: str,
    session: LenientClientSession,
    openai_client: OpenAI,
    state: ResumeState,
    emit: Emit,
    *,
    resume_file_id: str | None = None,
) -> None:
    """Handle a single chat turn for the Resume Optimizer.

    `resume_file_id` is supplied by the chat route when the user attached a
    PDF to this message. Once we've ingested the PDF, the text is stashed on
    `state["resume_text"]` so follow-up turns don't need to re-upload.
    """
    # 1) Pull PDF text if a new file was attached.
    if resume_file_id:
        loaded = load_upload(resume_file_id)
        if loaded is None:
            await emit(
                {
                    "type": "assistant_message",
                    "text": "I couldn't find the file you attached — try uploading it again?",
                }
            )
            return
        pdf_path, original_name = loaded
        await emit(
            {"type": "status", "message": f"Reading {original_name}…"}
        )
        state["resume_text"] = _parse_resume_pdf(pdf_path)
        state["resume_name"] = original_name
        if not state["resume_text"]:
            await emit(
                {
                    "type": "assistant_message",
                    "text": (
                        "I couldn't extract any text from that PDF — it might be a "
                        "scanned image rather than a text-based document. Can you "
                        "upload a text-based PDF instead?"
                    ),
                }
            )
            state["resume_text"] = None
            return

    # 2) Understand title + company.
    await emit({"type": "status", "message": "Understanding your request…"})
    prior = state.get("intent")
    intent = _extract_intent(openai_client, user_message, prior)
    state["intent"] = {
        "job_title": intent.job_title,
        "company": intent.company,
        "location": intent.location,
    }

    if not state.get("resume_text"):
        state["awaiting"] = "resume"
        await emit(
            {
                "type": "assistant_message",
                "text": (
                    "Upload your resume as a PDF using the paperclip button and "
                    "tell me the job title and company you're targeting (e.g. "
                    "*Senior Backend Engineer at Stripe*)."
                ),
            }
        )
        return
    if not intent.job_title:
        state["awaiting"] = "job_title"
        await emit(
            {
                "type": "assistant_message",
                "text": "What's the job title you're targeting?",
            }
        )
        return
    if not intent.company:
        state["awaiting"] = "company"
        await emit(
            {
                "type": "assistant_message",
                "text": f"Got it — **{intent.job_title}**. Which company?",
            }
        )
        return

    state["awaiting"] = None
    await emit(
        {
            "type": "intent",
            "niche": intent.job_title,
            "location": intent.company,
            "count": None,
            "hashtags": [intent.location] if intent.location else [],
        }
    )
    await emit(
        {
            "type": "assistant_message",
            "text": (
                f"Optimizing **{state.get('resume_name') or 'your resume'}** for "
                f"**{intent.job_title}** at **{intent.company}**. I'll scrape the "
                f"LinkedIn job post first, then tailor the resume."
            ),
        }
    )

    # 3) Scrape LinkedIn for the job description.
    await emit({"type": "status", "message": "Searching LinkedIn jobs…"})
    job_description = await _scrape_linkedin_job(session, intent, emit)
    if not job_description:
        await emit(
            {
                "type": "assistant_message",
                "text": (
                    f"I couldn't find a matching LinkedIn listing for "
                    f"*{intent.job_title}* at *{intent.company}*. I'll optimize "
                    "against a generic version of that role instead."
                ),
            }
        )
        job_description = (
            f"Target role: {intent.job_title} at {intent.company}. "
            "No specific listing found — use general best practices for this role."
        )

    # 4) Rewrite the resume.
    await emit({"type": "status", "message": "Rewriting your resume with GPT…"})
    markdown, summary = _rewrite_resume(
        openai_client, state["resume_text"], job_description, intent
    )
    if not markdown:
        await emit(
            {
                "type": "error",
                "message": "The model returned an empty resume. Try again?",
            }
        )
        return

    # 5) Save outputs + emit result event.
    md_name, docx_name = _save_outputs(markdown, intent)
    await emit(
        {
            "type": "resume_result",
            "job_title": intent.job_title,
            "company": intent.company,
            "markdown": markdown,
            "md_url": f"/api/downloads/{md_name}",
            "docx_url": f"/api/downloads/{docx_name}",
            "summary": summary,
        }
    )
    await emit(
        {
            "type": "assistant_message",
            "text": (
                f"Done! Your resume has been tailored for **{intent.job_title}** "
                f"at **{intent.company}**. Download the markdown or editable "
                ".docx on the right. Want another version for a different role? "
                "Just tell me the new title and company."
            ),
        }
    )
    # Clear intent so the next turn is treated as a fresh request, but keep
    # the resume text around so the user doesn't have to re-upload.
    state["intent"] = None
